from telegram import Update, ReplyKeyboardMarkup, KeyboardButton
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from telegram.constants import ChatAction
from docx import Document
import os

TOKEN = '8159286438:AAEhG_NgnX_NHHcRw4t9T6251-cHFyW_quo'

# Функция для старта
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    keyboard = [
        [KeyboardButton("ИП"), KeyboardButton("Самозанятый")],
    ]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)

    await update.message.reply_text(
        "Привет! Пожалуйста, выберите, с кем мы работаем: ИП или Самозанятый.",
        reply_markup=reply_markup,
    )

# Обработчик для выбора "Самозанятый"
async def handle_self_employed(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Вы выбрали Самозанятый. Пожалуйста, отправьте файл 'карточку' в формате .docx."
    )
    context.user_data["mode"] = "self_employed"

# Обработчик для выбора "ИП"
async def handle_individual_entrepreneur(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Вы выбрали ИП. Пожалуйста, отправьте файл 'карточку' в формате .docx."
    )
    context.user_data["mode"] = "individual_entrepreneur"

# Обработчик файла
async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    mode = context.user_data.get("mode")
    if mode not in ["self_employed", "individual_entrepreneur"]:
        await update.message.reply_text("Функция пока недоступна, выберите иной вариант.")
        return

    file = update.message.document
    if file.mime_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        await update.message.reply_text("Пожалуйста, отправьте файл в формате .docx.")
        return

    await update.message.reply_text("Файл получен! Идёт обработка...")
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action=ChatAction.TYPING)

    file_id = file.file_id
    tg_file = await context.bot.get_file(file_id)
    file_name = file.file_name
    file_path = os.path.join("/home/darkking/hr-telegram-bot/", file_name)

    context.user_data["file_name"] = file_name
    context.user_data["file_path"] = file_path

    await tg_file.download_to_drive(file_path)

    try:
        card_data = extract_data_from_docx(file_path)
        context.user_data["card_data"] = card_data

        if mode == "self_employed":
            await update.message.reply_text(
                "Пожалуйста, введите недостающие данные, каждое значение с новой строки:\n"
                "к/с\n"
                "Номер и дата договора\n"
                "Дата постановки на учет\n"
                "Справка о постановке\n"
                "Услуга исполнителя\n"
                "Стоимость\n"
                "Актуальная дата договора\n"
                "Сроки оказания\n"
                "Оплата услуг"
            )
        elif mode == "individual_entrepreneur":
            await update.message.reply_text(
                "Пожалуйста, введите недостающие данные для ИП, каждое значение с новой строки:\n"
                "к/с\n"
                "Номер договора\n"
                "Дата заключения договора\n"
                "Наименование услуги\n"
                "Стоимость услуги\n"
                "Срок оказания услуг\n"
                "Дополнительные условия (если есть)"
            )
        context.user_data["waiting_for_input"] = True
    except Exception as e:
        await update.message.reply_text(f"Произошла ошибка: {e}")

# Обработчик ввода недостающих данных
async def collect_missing_data(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if context.user_data.get("waiting_for_input"):
        user_message = update.message.text.split("\n")
        mode = context.user_data.get("mode")

        expected_count = 9 if mode == "self_employed" else 7
        if len(user_message) != expected_count:
            await update.message.reply_text(
                f"Ошибка: необходимо ввести {expected_count} значений, каждое с новой строки.\nПопробуйте снова."
            )
            return

        if mode == "self_employed":
            context.user_data["additional_data"] = {
                "к/с": user_message[0],
                "НОМЕР И ДАТА ДОГОВОРА": user_message[1],
                "ДАТА ПОСТАНОВКИ НА УЧЕТ": user_message[2],
                "СПРАВКА О ПОСТАНОВКЕ": user_message[3],
                "УСЛУГА ИСПОЛНИТЕЛЯ": user_message[4],
                "СТОИМОСТЬ": user_message[5],
                "АКТУАЛЬНАЯ ДАТА": user_message[6],
                "СРОКИ ОКАЗАНИЯ": user_message[7],
                "ОПЛАТА УСЛУГ": user_message[8],
            }
        elif mode == "individual_entrepreneur":
            context.user_data["additional_data"] = {
                "к/с": user_message[0],
                "НОМЕР ДОГОВОРА": user_message[1],
                "ДАТА ЗАКЛЮЧЕНИЯ": user_message[2],
                "НАИМЕНОВАНИЕ УСЛУГИ": user_message[3],
                "СТОИМОСТЬ УСЛУГИ": user_message[4],
                "СРОК ОКАЗАНИЯ УСЛУГ": user_message[5],
                "ДОПОЛНИТЕЛЬНЫЕ УСЛОВИЯ": user_message[6],
            }

        context.user_data["waiting_for_input"] = False

        await process_and_send_files(update, context)
    else:
        await update.message.reply_text("Я вас не понял. Пожалуйста, сначала отправьте файл.")

# Функция для обработки и отправки файлов
async def process_and_send_files(update, context):
    card_data = context.user_data["card_data"]
    additional_data = context.user_data["additional_data"]

    replacement_dict = {**card_data, **additional_data}
    file_name = context.user_data["file_name"]

    input_file1 = "/home/darkking/hr-telegram-bot/Акт_ИП.docx" if context.user_data["mode"] == "individual_entrepreneur" else "/home/darkking/hr-telegram-bot/Акт_СЗ.docx"
    output_file1 = os.path.join("/home/darkking/hr-telegram-bot/", f"Акт_{file_name}.docx")
    input_file2 = "/home/darkking/hr-telegram-bot/Договор_ИП.docx" if context.user_data["mode"] == "individual_entrepreneur" else "/home/darkking/hr-telegram-bot/Договор_СЗ.docx"
    output_file2 = os.path.join("/home/darkking/hr-telegram-bot/", f"Договор_{file_name}.docx")

    replace_highlighted_text(input_file1, output_file1, replacement_dict)
    replace_highlighted_text(input_file2, output_file2, replacement_dict)

    await context.bot.send_document(chat_id=update.effective_chat.id, document=open(output_file1, "rb"))
    await context.bot.send_document(chat_id=update.effective_chat.id, document=open(output_file2, "rb"))

    await update.message.reply_text("Файлы успешно обработаны и отправлены!\nВозвращаюсь в главное меню.")
    await start(update, context)

# Считывание данных из файла .docx
def extract_data_from_docx(file_path):
    doc = Document(file_path)
    data = {}
    for paragraph in doc.paragraphs:
        line = paragraph.text.strip()
        if ":" in line:
            key, value = map(str.strip, line.split(":", 1))
            if key and value:
                data[key] = value
    return data

# Замена текста в документе
def replace_highlighted_text(input_path, output_path, replacements, highlight_color=7):
    doc = Document(input_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color == highlight_color:
                for key, value in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        run.font.highlight_color = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.highlight_color == highlight_color:
                            for key, value in replacements.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                                    run.font.highlight_color = None
    doc.save(output_path)

# Запуск бота
def main():
    app = Application.builder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))

    print("Бот запущен!")
    app.run_polling()

if __name__ == "__main__":
    main()
