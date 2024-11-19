TOKEN = '8159286438:AAEhG_NgnX_NHHcRw4t9T6251-cHFyW_quo'

from docx import Document

def extract_data_from_docx(file_path):
    """
    Извлекает данные из документа, где каждая строка содержит формат 'Ключ: Значение'.
    """
    doc = Document(file_path)
    data = {}
    for paragraph in doc.paragraphs:
        line = paragraph.text.strip()
        if ":" in line:
            key, value = map(str.strip, line.split(":", 1))
            if key and value:
                data[key] = value
    return data

def replace_highlighted_text(input_path, output_path, replacements, highlight_color=7):
    """
    Заменяет текстовые метки в документе, выделенные определённым цветом, на значения из словаря.
    """
    doc = Document(input_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color == highlight_color:
                for key, value in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        run.font.highlight_color = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.highlight_color == highlight_color:
                            for key, value in replacements.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                                    run.font.highlight_color = None
    doc.save(output_path)
    print(f"Файл успешно сохранён: {output_path}")

card_file = "/home/darkking/hr-telegram-bot/Карточка самозанятого.docx"
input_file1 = "/home/darkking/hr-telegram-bot/Акт_СЗ.docx"
output_file1 = "/home/darkking/hr-telegram-bot/Акт_СЗ_готовый.docx"
input_file2 = "/home/darkking/hr-telegram-bot/Договор_СЗ.docx"
output_file2 = "/home/darkking/hr-telegram-bot/Договор_СЗ_готовый.docx"

card_data = extract_data_from_docx(card_file)

fio_full = card_data.get("ФИО(полностью)", "")
fio_short = card_data.get("ФИО(сокращено)", "")
birth_date = card_data.get("Дата рождения", "")
address = card_data.get("Адрес регистрации", "")
passport = card_data.get("Паспорт(серия номер)", "")
bank_name = card_data.get("Наименование банка", "")
account_number = card_data.get("Номер счета", "")
bik = card_data.get("БИК", "")
department_code = card_data.get("Код подразделения", "")
inn = card_data.get("ИНН", "")

k_s = "0000000"
number_date_agreement = "123 01.01.01"
registration_date = "02.02.02"
certificate = "99999"
work_type = "Тестоваяя услуга"
cost = "111"
currrent_data = "01.01.01"
terms_of_provision = "0 февраля 0000 г. по 00 февраля 0000 г"
services_payment = "не позднее 27 февраля 2024 года"

replacement_dict = {
    "ФИОп": fio_full,
    "ФИО": fio_short,
    "ПАСПОРТ:": passport,
    "ДАТА РОЖДЕНИЯ:": birth_date,
    "АДРЕС:": address,
    "ИНН:": inn,
    "НАИМЕНОВАНИЕ БАНКА:": bank_name,
    "НОМЕР СЧЕТА:": account_number,
    "БИК": bik,
    "к/с": k_s,
    "КОД ПОДРАЗДЕЛЕНИЯ:": department_code,
    "ФИО": fio_short,
    "НОМЕР И ДАТА ДОГОВОРА": number_date_agreement,
    "ДАТА ПОСТАНОВКИ НА УЧЕТ": registration_date,
    "СПРАВКА О ПОСТАНОВКЕ": certificate,
    "УСЛУГА ИСПОЛНИТЕЛЯ": work_type,
    "СТОИМОСТЬ": cost,
    "АКТУАЛЬНАЯ ДАТА": currrent_data,
    "СРОКИ ОКАЗАНИЯ": terms_of_provision,
    "ОПЛАТА УСЛУГ": services_payment
}

replace_highlighted_text(input_file1, output_file1, replacement_dict)
replace_highlighted_text(input_file2, output_file2, replacement_dict)

