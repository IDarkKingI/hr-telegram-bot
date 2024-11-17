TOKEN = '8159286438:AAEhG_NgnX_NHHcRw4t9T6251-cHFyW_quo'

from docx import Document

def update_replacement_dict(base_dict, card_file_path):
    """
    Обновляет словарь замен на основе данных из файла "Карточка самозанятого".
    :param base_dict: исходный словарь с заменами
    :param card_file_path: путь к файлу карточки
    :return: обновлённый словарь
    """
    doc = Document(card_file_path)
    
    # Новый словарь с данными из карточки
    updated_dict = base_dict.copy()

    # Чтение строк из документа
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if "ФИО(полностью):" in text:
            updated_dict["ФИОп"] = text.split("ФИО(полностью):")[1].strip()
        elif "Дата рождения:" in text:
            updated_dict["ДАТА РОЖДЕНИЯ:"] = text.split("Дата рождения:")[1].strip()
        elif "Адрес регистрации:" in text:
            updated_dict["АДРЕС:"] = text.split("Адрес регистрации:")[1].strip()
        elif "Паспорт: серия номер:" in text:
            updated_dict["ПАСПОРТ:"] = text.split("Паспорт(серия номер):")[1].strip()
        elif "код подразделения:" in text:
            updated_dict["КОД ПОДРАЗДЕЛЕНИЯ:"] = text.split("код подразделения:")[1].strip()
        elif "ИНН:" in text:
            updated_dict["ИНН:"] = text.split("ИНН:")[1].strip()
        elif "Наименование банка:" in text:
            updated_dict["НАИМЕНОВАНИЕ БАНКА:"] = text.split("Наименование банка:")[1].strip()
        elif "БИК:" in text:
            updated_dict["БИК"] = text.split("БИК:")[1].strip()
        elif "Номер счета:" in text:
            updated_dict["НОМЕР СЧЕТА:"] = text.split("Номер счета:")[1].strip()

    return updated_dict

def replace_highlighted_text(input_path, output_path, replacements, highlight_color=7):
    """
    Заменяет текстовые метки в документе, выделенные определённым цветом, на значения из словаря.
    """
    doc = Document(input_path)

    # Обрабатываем абзацы
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color == highlight_color:  # Проверяем выделение
                for key, value in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        run.font.highlight_color = None  # Убираем выделение

    # Обрабатываем таблицы
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.highlight_color == highlight_color:  # Проверяем выделение
                            for key, value in replacements.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                                    run.font.highlight_color = None  # Убираем выделение

    # Сохраняем измененный файл
    doc.save(output_path)
    print(f"Файл успешно сохранён: {output_path}")

# Путь к файлу карточки
card_file = "/home/darkking/hr-telegram-bot/Карточка самозанятого.docx"

# Путь к исходному и выходному файлу
input_file = "/home/darkking/hr-telegram-bot/Акт_СЗ.docx"
output_file = "/home/darkking/hr-telegram-bot/Акт_СЗ_готовый.docx"

# Исходный словарь замен
replacement_dict = {
    "ФИОп": "Иванов Иван Иванович",
    "ПАСПОРТ:": "1234 567890",
    "ДАТА РОЖДЕНИЯ:": "01.01.1990",
    "АДРЕС:": "г. Москва, ул. Пушкина, д. 10",
    "ИНН:": "123456789012",
    "НАИМЕНОВАНИЕ БАНКА:": "123456789012",
    "НОМЕР СЧЕТА:": "99999999999999999",
    "БИК": "3333333333",
    "к/с": "111111111",
    "КОД ПОДРАЗДЕЛЕНИЯ:": "123-456",
    "ФИО": "Иванов И.И.",
    "НОМЕР И ДАТА ДОГОВОРА": "№ 123/456 от 01.01.2024",
    "ДАТА ПОСТАНОВКИ НА УЧЕТ": "01.01.2020",
    "СПРАВКА О ПОСТАНОВКЕ": "№ 789456",
    "УСЛУГА ИСПОЛНИТЕЛЯ": "Разработка графического дизайна",
    "СТОИМОСТЬ": "10 000 (Десять тысяч) рублей 00 копеек",
    "АКТУАЛЬНАЯ ДАТА": "22 февраля 2024 года"
}

# Обновляем словарь данными из карточки
replacement_dict = update_replacement_dict(replacement_dict, card_file)

print("\nОбновленный словарь замен:")
for key, value in replacement_dict.items():
    print(f"{key}: {value}")

# Выполняем замену
replace_highlighted_text(input_file, output_file, replacement_dict)
