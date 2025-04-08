from telegram import Update, ReplyKeyboardMarkup, KeyboardButton
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from telegram.constants import ChatAction
from docx import Document
from num2words import num2words
import os


TOKEN = '8159286438:AAEhG_NgnX_NHHcRw4t9T6251-cHFyW_quo'

def convert_price_to_text(amount):
    try:
        if amount < 0:
            raise ValueError("Введите положительное число.")

        rubles = int(amount)
        kopecks = int(round((amount - rubles) * 100))

        rubles_words = f"{num2words(rubles, lang='ru')} {get_currency_form(rubles, ['рубль', 'рубля', 'рублей'])}"
        kopecks_words = f"{kopecks:02d} {get_currency_form(kopecks, ['копейка', 'копейки', 'копеек'])}"

        return f"{rubles} ({rubles_words}), {kopecks_words}"
    except ValueError as e:
        return f"Ошибка: {e}"

def get_currency_form(number, forms):
    if 11 <= number % 100 <= 19:
        return forms[2]
    elif number % 10 == 1:
        return forms[0]
    elif 2 <= number % 10 <= 4:
        return forms[1]
    else:
        return forms[2]


async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    keyboard = [
        [KeyboardButton("ИП"), KeyboardButton("Самозанятый")],
    ]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)

    await update.message.reply_text(
        "Привет! Пожалуйста, выберите, с кем мы работаем: ИП или Самозанятый.",
        reply_markup=reply_markup,
    )

async def handle_self_employed(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Вы выбрали Самозанятый. Пожалуйста, отправьте файл 'карточку' Самозанятого в формате .docx."
    )
    context.user_data["mode"] = "self_employed"

async def handle_individual_entrepreneur(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Вы выбрали ИП. Пожалуйста, отправьте файл 'карточку' ИП в формате .docx."
    )
    context.user_data["mode"] = "individual_entrepreneur"

async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    mode = context.user_data.get("mode")
    if mode not in ["self_employed", "individual_entrepreneur"]:
        await update.message.reply_text("Функция пока недоступна, выберите иной вариант.")
        return

    file = update.message.document
    if file.mime_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        await update.message.reply_text("Пожалуйста, отправьте файл в формате .docx.")
        return

    await update.message.reply_text("Файл получен! Идёт обработка...")
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action=ChatAction.TYPING)

    file_id = file.file_id
    tg_file = await context.bot.get_file(file_id)
    file_name = file.file_name
    file_path = os.path.join("/mnt/hr-telegram-bot/", file_name)

    context.user_data["file_name"] = file_name
    context.user_data["file_path"] = file_path

    await tg_file.download_to_drive(file_path)

    try:
        card_data = extract_data_from_docx(file_path)
        context.user_data["card_data"] = card_data

        if mode == "self_employed":
            await update.message.reply_text(
                "Пожалуйста, введите недостающие данные, каждое значение с новой строки:\n"
                "Услуга исполнителя\n"
                "Стоимость услуги\n"
                "Крайний срок оплаты\n"
                "Сроки оказания услуг"
            )
        elif mode == "individual_entrepreneur":
            await update.message.reply_text(
                "Пожалуйста, введите недостающие данные для ИП, каждое значение с новой строки:\n"
                "Наименование услуги\n"
                "Стоимость услуги\n"
                "Срок оказания услуг\n"
                "Формат результата"
            )
        context.user_data["waiting_for_input"] = True
    except Exception as e:
        await update.message.reply_text(f"Произошла ошибка: {e}")

async def collect_missing_data(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if context.user_data.get("waiting_for_input"):
        user_message = update.message.text.split("\n")
        mode = context.user_data.get("mode")
        expected_count = 4 if mode == "self_employed" else 4
        if len(user_message) != expected_count:
            await update.message.reply_text(
                f"Ошибка: необходимо ввести {expected_count} значений, каждое с новой строки.\nПопробуйте снова."
            )
            return
        
        if mode == "self_employed":
            context.user_data["additional_data"] = {
                "УСЛУГА ИСПОЛНИТЕЛЯ": user_message[0],
                "СТОИМОСТЬ": user_message[1],
                "СПОСОБ ОПЛАТЫ": user_message[2],
                "СРОКИ ОКАЗАНИЯ": user_message[3],
            }
        
        elif mode == "individual_entrepreneur":
            context.user_data["additional_data"] = {
                "УСЛУГА ИСПОЛНИТЕЛЯ": user_message[0],
                "СТОИМОСТЬ": user_message[1],
                "СРОКИ ОКАЗАНИЯ": user_message[2],
                "ФОРМАТ РЕЗУЛЬТАТА": user_message[3],
            }

        context.user_data["waiting_for_input"] = False

        await process_and_send_files(update, context)
    else:
        await update.message.reply_text("Я вас не понял. Пожалуйста, сначала отправьте файл.")

async def process_and_send_files(update, context):
    card_data = context.user_data["card_data"]
    additional_data = context.user_data["additional_data"]

    cost = additional_data.get("СТОИМОСТЬ", "").replace(",", ".").strip()
    try:
        cost_float = float(cost)
        cost_text = convert_price_to_text(cost_float)
        cost_with_text = f"{cost_text}"
    except ValueError:
        cost_with_text = cost 
    
    replacement_dict = {
        "сокрФИО": card_data.get("ФИО(сокращенно)", ""),
        "ФИО": card_data.get("ФИО(полностью)", ""),
        "ИПполностью": card_data.get("ИП(полностью)", ""),
        "ИПсокращенно": card_data.get("ИП(полностью)", ""),
        "ПАСПОРТ:": card_data.get("Паспорт(серия номер)", ""),
        "ДАТА РОЖДЕНИЯ:": card_data.get("Дата рождения", ""),
        "АДРЕС:": card_data.get("Адрес регистрации", ""),
        "ИНН:": card_data.get("ИНН", ""),
        "НАИМЕНОВАНИЕ БАНКА:": card_data.get("Наименование банка", ""),
        "НОМЕР СЧЕТА:": card_data.get("Номер счета", ""),
        "БИК": card_data.get("БИК", ""),
        "ОГРНИП": card_data.get("ОГРНИП", ""),
        "ВЫДАН": card_data.get("Выдан", ""),
        "К/С": card_data.get("Корреспондентский счет", ""),
        "Р/С": card_data.get("Расчетный счет", ""),
        "к/с": card_data.get("к/с", ""),
        "ДАТА ПОСТАНОВКИ НА УЧЕТ": card_data.get("Дата постановки на учет", ""),
        "СПРАВКА О ПОСТАНОВКЕ": card_data.get("Справка о постановке на учет", ""),
        "УСЛУГА ИСПОЛНИТЕЛЯ": additional_data.get("УСЛУГА ИСПОЛНИТЕЛЯ", ""),
        "ФОРМАТ РЕЗУЛЬТАТА": additional_data.get("ФОРМАТ РЕЗУЛЬТАТА", ""),
        "ОПЛАТА УСЛУГ": additional_data.get("СПОСОБ ОПЛАТЫ", ""),
        "СТОИМОСТЬ": cost_with_text, 
        "СРОКИ ОКАЗАНИЯ": additional_data.get("СРОКИ ОКАЗАНИЯ", ""),
    }

    file_name = context.user_data["file_name"]

    input_file1 = "/mnt/hr-telegram-bot/Акт_ИП.docx" if context.user_data["mode"] == "individual_entrepreneur" else "/mnt/hr-telegram-bot/Акт_СЗ.docx"
    output_file1 = os.path.join("/mnt/hr-telegram-bot/", f"Акт_{file_name}")
    input_file2 = "/mnt/hr-telegram-bot/Договор_ИП.docx" if context.user_data["mode"] == "individual_entrepreneur" else "/mnt/hr-telegram-bot/Договор_СЗ.docx"
    output_file2 = os.path.join("/mnt/hr-telegram-bot/", f"Договор_{file_name}")

    replace_highlighted_text(input_file1, output_file1, replacement_dict)
    replace_highlighted_text(input_file2, output_file2, replacement_dict)

    await context.bot.send_document(chat_id=update.effective_chat.id, document=open(output_file1, "rb"))
    await context.bot.send_document(chat_id=update.effective_chat.id, document=open(output_file2, "rb"))

    await update.message.reply_text("Файлы успешно обработаны и отправлены!\nВозвращаюсь в главное меню.")
    await start(update, context)

def extract_data_from_docx(file_path):
    doc = Document(file_path)
    data = {}
    for paragraph in doc.paragraphs:
        line = paragraph.text.strip()
        if ":" in line:
            key, value = map(str.strip, line.split(":", 1))
            if key and value:
                data[key] = value
    return data

def replace_highlighted_text(input_path, output_path, replacements, highlight_color=7):
    doc = Document(input_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color == highlight_color:
                for key, value in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        run.font.highlight_color = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.highlight_color == highlight_color:
                            for key, value in replacements.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                                    run.font.highlight_color = None
    doc.save(output_path)

def main():
    app = Application.builder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Regex("^Самозанятый$"), handle_self_employed))
    app.add_handler(MessageHandler(filters.Regex("^ИП$"), handle_individual_entrepreneur))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_file))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, collect_missing_data))

    print("Бот запущен!")
    app.run_polling()

if __name__ == "__main__":
    main()
