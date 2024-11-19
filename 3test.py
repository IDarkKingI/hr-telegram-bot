from telegram import Update, ReplyKeyboardMarkup, KeyboardButton
from telegram.ext import Application, CommandHandler, MessageHandler, filters, ContextTypes
from telegram.constants import ChatAction
from docx import Document
import os

TOKEN = '8159286438:AAEhG_NgnX_NHHcRw4t9T6251-cHFyW_quo'


# Функция для старта
async def start(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    keyboard = [
        [KeyboardButton("ИП"), KeyboardButton("Самозанятый")],
    ]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True)

    await update.message.reply_text(
        "Привет! Пожалуйста, выберите, с кем мы работаем: ИП или самозанятый.",
        reply_markup=reply_markup,
    )


# Обработчик для выбора "самозанятый"
async def handle_self_employed(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    await update.message.reply_text(
        "Вы выбрали Самозанятый. Пожалуйста, отправьте файл 'карточку' в формате .docx."
    )
    context.user_data["mode"] = "self_employed"  # Устанавливаем режим для дальнейшей логики


# Обработчик файла
async def handle_file(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if context.user_data.get("mode") != "self_employed":
        await update.message.reply_text("Функция пока недоступна, выберите иной вариант")
        return

    file = update.message.document
    if file.mime_type != "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
        await update.message.reply_text("Пожалуйста, отправьте файл в формате .docx.")
        return

    await update.message.reply_text("Файл получен! Идёт обработка...")
    await context.bot.send_chat_action(chat_id=update.effective_chat.id, action=ChatAction.TYPING)

    file_id = file.file_id
    tg_file = await context.bot.get_file(file_id)
    file_path = f"{file.file_id}.docx"

    # Скачать файл
    await tg_file.download_to_drive(file_path)

    try:
        # Считать данные из файла
        card_data = extract_data_from_docx(file_path)
        context.user_data["card_data"] = card_data

        # Запрос недостающих данных у пользователя
        await update.message.reply_text(
            "Пожалуйста, введите недостающие данные, каждое значение с новой строки:\n"
            "к/с\n"
            "Номер и дата договора\n"
            "Дата постановки на учет\n"
            "Справка о постановке\n"
            "Услуга исполнителя\n"
            "Стоимость\n"
            "Актуальная дата договора\n"
            "Сроки оказания\n"
            "Оплата услуг"
        )
        context.user_data["waiting_for_input"] = True
    except Exception as e:
        await update.message.reply_text(f"Произошла ошибка: {e}")
    finally:
        if os.path.exists(file_path):
            os.remove(file_path)


# Обработчик ввода недостающих данных
async def collect_missing_data(update: Update, context: ContextTypes.DEFAULT_TYPE) -> None:
    if context.user_data.get("waiting_for_input"):
        user_message = update.message.text.split("\n")

        if len(user_message) != 9:
            await update.message.reply_text(
                "Ошибка: необходимо ввести 10 значений, каждое с новой строки.\nПопробуйте снова."
            )
            return

        # Привязка данных к переменным
        context.user_data["additional_data"] = {
            "к/с": user_message[0],
            "НОМЕР И ДАТА ДОГОВОРА": user_message[1],
            "ДАТА ПОСТАНОВКИ НА УЧЕТ": user_message[2],
            "СПРАВКА О ПОСТАНОВКЕ": user_message[3],
            "УСЛУГА ИСПОЛНИТЕЛЯ": user_message[4],
            "СТОИМОСТЬ": user_message[5],
            "АКТУАЛЬНАЯ ДАТА": user_message[6],
            "СРОКИ ОКАЗАНИЯ": user_message[7],
            "ОПЛАТА УСЛУГ": user_message[8],
        }

        context.user_data["waiting_for_input"] = False

        # Обработка файла
        await process_and_send_files(update, context)
    else:
        await update.message.reply_text("Я вас не понял. Пожалуйста, сначала отправьте файл.")


# Функция для обработки и отправки файлов
async def process_and_send_files(update, context):
    card_data = context.user_data["card_data"]
    additional_data = context.user_data["additional_data"]

    replacement_dict = {
        "ФИОп": card_data.get("ФИО(полностью)", ""),
        "ФИО": card_data.get("ФИО(сокращенно)", ""),
        "ПАСПОРТ:": card_data.get("Паспорт(серия номер)", ""),
        "ДАТА РОЖДЕНИЯ:": card_data.get("Дата рождения", ""),
        "АДРЕС:": card_data.get("Адрес регистрации", ""),
        "ИНН:": card_data.get("ИНН", ""),
        "НАИМЕНОВАНИЕ БАНКА:": card_data.get("Наименование банка", ""),
        "НОМЕР СЧЕТА:": card_data.get("Номер счета", ""),
        "БИК": card_data.get("БИК", ""),
        "КОД ПОДРАЗДЕЛЕНИЯ:": card_data.get("Код подразделения", ""),
        "к/с": additional_data["к/с"],
        "НОМЕР И ДАТА ДОГОВОРА": additional_data["НОМЕР И ДАТА ДОГОВОРА"],
        "ДАТА ПОСТАНОВКИ НА УЧЕТ": additional_data["ДАТА ПОСТАНОВКИ НА УЧЕТ"],
        "СПРАВКА О ПОСТАНОВКЕ": additional_data["СПРАВКА О ПОСТАНОВКЕ"],
        "УСЛУГА ИСПОЛНИТЕЛЯ": additional_data["УСЛУГА ИСПОЛНИТЕЛЯ"],
        "СТОИМОСТЬ": additional_data["СТОИМОСТЬ"],
        "АКТУАЛЬНАЯ ДАТА": additional_data["АКТУАЛЬНАЯ ДАТА"],
        "СРОКИ ОКАЗАНИЯ": additional_data["СРОКИ ОКАЗАНИЯ"],
        "ОПЛАТА УСЛУГ": additional_data["ОПЛАТА УСЛУГ"],
    }

    input_file1 = "/home/darkking/hr-telegram-bot/Акт_СЗ.docx"
    output_file1 = "/home/darkking/hr-telegram-bot/Акт_СЗ_готовый.docx"
    input_file2 = "/home/darkking/hr-telegram-bot/Договор_СЗ.docx"
    output_file2 = "/home/darkking/hr-telegram-bot/Договор_СЗ_готовый.docx"

    replace_highlighted_text(input_file1, output_file1, replacement_dict)
    replace_highlighted_text(input_file2, output_file2, replacement_dict)

    await context.bot.send_document(chat_id=update.effective_chat.id, document=open(output_file1, "rb"))
    await context.bot.send_document(chat_id=update.effective_chat.id, document=open(output_file2, "rb"))

    # Вернуться в главное меню
    await update.message.reply_text("Файлы успешно обработаны и отправлены!\nВозвращаюсь в главное меню.")
    await start(update, context)


# Считывание данных из файла .docx
def extract_data_from_docx(file_path):
    doc = Document(file_path)
    data = {}
    for paragraph in doc.paragraphs:
        line = paragraph.text.strip()
        if ":" in line:
            key, value = map(str.strip, line.split(":", 1))
            if key and value:
                data[key] = value
    return data


# Замена текста в документе
def replace_highlighted_text(input_path, output_path, replacements, highlight_color=7):
    doc = Document(input_path)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run.font.highlight_color == highlight_color:
                for key, value in replacements.items():
                    if key in run.text:
                        run.text = run.text.replace(key, value)
                        run.font.highlight_color = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.font.highlight_color == highlight_color:
                            for key, value in replacements.items():
                                if key in run.text:
                                    run.text = run.text.replace(key, value)
                                    run.font.highlight_color = None
    doc.save(output_path)


# Запуск бота
def main():
    app = Application.builder().token(TOKEN).build()

    app.add_handler(CommandHandler("start", start))
    app.add_handler(MessageHandler(filters.Regex("^Самозанятый$"), handle_self_employed))
    app.add_handler(MessageHandler(filters.Document.ALL, handle_file))
    app.add_handler(MessageHandler(filters.TEXT & ~filters.COMMAND, collect_missing_data))

    print("Бот запущен!")
    app.run_polling()


if __name__ == "__main__":
    main()
